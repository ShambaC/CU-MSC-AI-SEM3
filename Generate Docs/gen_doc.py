from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_assignment_docx():
    # Create a new Document
    doc = Document()

    # Title (Centered and Bold)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Assignment 13')
    run.bold = True
    run.font.size = Pt(14)  # Adjust font size if needed

    # Problem Statement (Bold)
    doc.add_heading('Problem Statement:', level=1)
    problem_statement = doc.add_paragraph('Perform the Traveling Salesman Problem (TSP) using a Genetic Algorithm (GA).')
    problem_statement.runs[0].bold = True

    # Description
    doc.add_heading('Description:', level=1)
    doc.add_paragraph(
        'The Traveling Salesman Problem (TSP) is a classic optimization problem where the goal is to find the shortest possible route that visits a set of cities exactly once and returns to the starting city. '
        'The Genetic Algorithm (GA) is a heuristic search algorithm inspired by the process of natural selection. It is used to find approximate solutions to optimization problems like TSP.'
    )
    doc.add_paragraph(
        'In this assignment, the GA is applied to solve the TSP. The algorithm involves creating a population of possible routes, encoding them, selecting the best routes based on fitness, and applying crossover and mutation to generate new routes. '
        'The process is repeated for a specified number of generations to find the optimal or near-optimal route.'
    )

    # Algorithm (Numbered Steps with Bold "Step n:")
    doc.add_heading('Algorithm:', level=1)
    algorithm_steps = [
        "Generate an initial population of random routes (permutations of cities).",
        "Encode each route into a binary or integer representation for easier manipulation.",
        "Calculate the fitness of each route based on the total distance traveled (shorter distances have higher fitness).",
        "Select parent routes for reproduction using a weighted probability based on their fitness scores.",
        "Perform crossover (e.g., Ordered Crossover) on the selected parents to produce offspring routes.",
        "Apply mutation to the offspring routes to introduce genetic diversity.",
        "Replace the least fit routes in the population with the new offspring routes.",
        "Repeat the selection, crossover, mutation, and replacement steps for a specified number of generations.",
        "Return the route with the highest fitness (shortest distance) as the solution."
    ]

    for i, step in enumerate(algorithm_steps, start=1):
        step_paragraph = doc.add_paragraph()
        step_run = step_paragraph.add_run(f'Step {i}: ')
        step_run.bold = True  # Make "Step n:" bold
        step_paragraph.add_run(step)  # Add the rest of the step text

    # Source Code (Stub)
    doc.add_heading('Source Code:', level=1)
    source_code_stub = """
import numpy as np
import click
import random

from itertools import permutations
from tqdm import tqdm, trange
from pprint import pp

def generate_population(num_cities: int, pop_size: int) -> list:
    # Generate initial population of routes
    ...

def generate_encode_list(num_cities: int) -> dict:
    # Generate encoding for each city
    ...

def encode(route: tuple[int], encode_list: dict) -> tuple[str]:
    # Encode a route into binary or integer representation
    ...

def decode(route: tuple[str], encode_list: dict) -> tuple[int]:
    # Decode a route back to city indices
    ...

def calc_fitness(route: tuple[str], distance_mat: np.ndarray, encode_list: dict, is_distance: bool=False) -> int:
    # Calculate the fitness of a route based on total distance
    ...

def selection(routes: dict) -> np.ndarray:
    # Select parents based on fitness scores
    ...

def crossover(parentA: tuple[str], parentB: tuple[str], num_cities: int) -> tuple[str]:
    # Perform crossover to generate offspring routes
    ...

def mutation(child: tuple[str], num_cities: int) -> tuple[str]:
    # Apply mutation to introduce genetic diversity
    ...

@click.command()
@click.option('--num', '-N', default=5, help='Number of cities')
@click.option('--pop_size', '-P', default=10, help='Population size')
@click.option('--gen_count', '-G', default=50, help='Number of generations')
@click.option('--mut_prob', '-M', default=0.2, help='Mutation probability')
@click.option('--seed', '-S', default=42, help='Seed for RNG')
def main(num, pop_size, gen_count, mut_prob, seed):
    # Main function to run the Genetic Algorithm for TSP
    ...

if __name__ == '__main__':
    main()
    """

    # Add the source code stub with monospace font and smaller font size
    code_paragraph = doc.add_paragraph()
    run = code_paragraph.add_run(source_code_stub)
    font = run.font
    font.name = 'Courier New'  # Monospace font
    font.size = Pt(10)  # Smaller font size for code

    # Output (Placeholder)
    doc.add_heading('Output:', level=1)
    doc.add_paragraph('[Placeholder for output image or graph]')

    # Report
    doc.add_heading('Report:', level=1)
    doc.add_heading('Performance of Genetic Algorithm for TSP:', level=2)
    doc.add_paragraph(
        '- The Genetic Algorithm (GA) is effective for solving the Traveling Salesman Problem (TSP) and other combinatorial optimization problems.\n'
        '- It can find near-optimal solutions even for large problem sizes, although it may not guarantee the global optimum.\n'
        '- The algorithm has a time complexity of O(n^2) per generation, making it suitable for medium-sized datasets.'
    )

    doc.add_heading('Drawbacks of Genetic Algorithm for TSP:', level=2)
    doc.add_paragraph(
        '- Scalability: Due to its high time complexity, GA is not suitable for very large datasets.\n'
        '- Sensitivity to Parameters: The performance of the algorithm can be highly dependent on parameters like population size, mutation rate, and crossover method.\n'
        '- No Guarantee of Optimality: GA may converge to a local optimum rather than the global optimum.'
    )

    doc.add_heading('Remedies:', level=2)
    doc.add_paragraph(
        '- For large datasets, consider using more scalable algorithms like simulated annealing or ant colony optimization.\n'
        '- Experiment with different parameter settings to improve the performance of the GA.\n'
        '- Use techniques like elitism or hybrid approaches to improve convergence and solution quality.'
    )

    # Save the document
    doc.save('Assignment_13_Genetic_Algorithm_TSP.docx')

# Run the function to generate the document
create_assignment_docx()